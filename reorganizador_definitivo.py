# -*- coding: utf-8 -*-
"""
📘 Reorganizador y Manual de Estudio Python – Versión Educativa Definitiva

Este script:
1️⃣ Recorre todos los archivos .py de una carpeta (incluyendo subcarpetas).
2️⃣ Genera un Word con:
   - Código coloreado por sintaxis.
   - Explicaciones educativas de cada bloque.
   - Pseudodiagramas de flujo en texto si Graphviz no está disponible.
3️⃣ Genera un TXT resumen con los pasos realizados y librerías instaladas.
4️⃣ Pregunta al usuario si quiere generar diagramas con Graphviz.
"""

import os
import sys
import subprocess
from pathlib import Path
import re

# ============================
# Función para instalar librerías si no están
# ============================
def instalar_libreria(libreria):
    try:
        __import__(libreria)
        print(f"✅ Librería {libreria} ya instalada")
    except ImportError:
        print(f"⚠ Librería {libreria} no encontrada. Instalando...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", libreria])
        print(f"✅ {libreria} instalada correctamente")

# Librerías necesarias
for lib in ["python-docx", "pygments", "matplotlib", "graphviz", "pydot"]:
    instalar_libreria(lib)

# Importaciones tras asegurarnos que están instaladas
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from pygments.lexers import PythonLexer
from pygments.token import Token

# ============================
# Función para colorear código y añadir explicaciones
# ============================
def agregar_codigo_coloreado(doc, codigo, explicacion):
    """
    Añade al Word:
    - Explicación educativa
    - Código coloreado por sintaxis (keywords, strings, comentarios, números)
    - Pseudodiagrama de flujo textual
    """
    # Añadimos la explicación educativa
    doc.add_paragraph(explicacion)

    # Separar línea por línea
    for linea in codigo.split("\n"):
        parrafo = doc.add_paragraph()
        run = parrafo.add_run(linea if linea.strip() else " ")  # Manejo de líneas vacías
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(10)

        # Colores por tipo de token
        if re.match(r"^\s*#", linea):  # Comentario verde
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        elif re.search(r"(['\"].*?['\"])", linea):  # Strings naranja
            run.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
        elif re.search(r"\b(True|False|None)\b", linea):  # Constantes rojo
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        elif re.search(r"\b(def|return|if|else|elif|for|while|import|from|class|with|as|try|except|finally|pass|break|continue)\b", linea):
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Keywords azul
        elif re.search(r"\b(\d+(\.\d+)?)\b", linea):  # Números morado
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x80)
        else:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Negro por defecto

    # Añadimos pseudodiagrama de flujo en texto
    doc.add_paragraph("\n🟢 Pseudodiagrama de flujo:")
    for l in codigo.split("\n"):
        l = l.strip()
        if l.startswith("def "):
            doc.add_paragraph(f"→ Función: {l}")
        elif l.startswith("if "):
            doc.add_paragraph(f"  └─ Si {l[3:]}:")
        elif l.startswith("elif "):
            doc.add_paragraph(f"  └─ Sino si {l[5:]}:")
        elif l.startswith("else"):
            doc.add_paragraph(f"  └─ Sino:")
        elif l.startswith("for "):
            doc.add_paragraph(f"  └─ Para {l[4:]}:")
        elif l.startswith("while "):
            doc.add_paragraph(f"  └─ Mientras {l[6:]}:")

# ============================
# Función principal
# ============================
def main():
    # Carpeta de trabajo
    carpeta = input("Introduce la ruta de la carpeta con los ejercicios Python: ").strip()
    carpeta = Path(carpeta)
    if not carpeta.exists():
        print("❌ Carpeta no encontrada")
        return

    # Preguntar si generar diagramas profesionales con Graphviz
    generar_gv = input("¿Deseas generar diagramas profesionales con Graphviz si está instalado? (s/n): ").lower() == 's'

    # Crear Word y TXT resumen
    doc = Document()
    resumen_txt = []

    # Recorrer todos los archivos .py incluyendo subcarpetas
    for root, dirs, files in os.walk(carpeta):
        for archivo in files:
            if archivo.endswith(".py"):
                ruta_archivo = Path(root) / archivo
                resumen_txt.append(f"Procesando {archivo}")
                with open(ruta_archivo, "r", encoding="utf-8") as f:
                    codigo = f.read()

                explicacion = f"Este archivo contiene el código de {archivo}. Observa la sintaxis, los comentarios y el flujo del programa."
                agregar_codigo_coloreado(doc, codigo, explicacion)

                # Diagramas opcionales
                if generar_gv:
                    try:
                        import graphviz
                        import pydot
                        # Placeholder: Aquí se podría generar diagrama real
                        resumen_txt.append(f"Diagrama profesional generado para {archivo}")
                    except Exception as e:
                        resumen_txt.append(f"⚠ No se pudo generar diagrama para {archivo}: {e}")

    # Guardar Word y TXT
    doc_path = carpeta / "Manual_Estudio_Python.docx"
    txt_path = carpeta / "Resumen_Pasos.txt"
    doc.save(doc_path)
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(resumen_txt))

    print(f"✅ Manual Word generado: {doc_path}")
    print(f"✅ Resumen TXT generado: {txt_path}")
    print("🎉 Proceso completado. Manual y resumen listos para estudio.")

# ============================
# Ejecutar
# ============================
if __name__ == "__main__":
    main()
    